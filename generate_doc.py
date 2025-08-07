import markdown
from weasyprint import HTML
import pathlib
from base import Resume
import os
import re
from docx import Document as DocxDocument
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
def setup_document_styles(document):
    """
    设置文档的全局和标题样式。
    """
    # --- 1. 定义基础颜色和字体 ---
    primary_color = RGBColor(79, 129, 189)  # 一种专业的蓝色
    font_name = '微软雅黑'  # 使用更现代的字体

    # --- 2. 设置默认正文样式 ('Normal') ---
    style = document.styles['Normal']
    font = style.font
    font.name = font_name
    font.size = Pt(10.5)
    # 设置中文字体
    rpr = style.element.rPr
    rpr.rFonts.set(qn('w:eastAsia'), font_name)

    # --- 3. 自定义标题样式 ---
    # 标题1 (H1)
    style_h1 = document.styles.add_style('CustomHeading1', 1)
    style_h1.base_style = document.styles['Heading 1']
    font_h1 = style_h1.font
    font_h1.name = font_name
    font_h1.size = Pt(18)
    font_h1.bold = True
    font_h1.color.rgb = primary_color
    style_h1.paragraph_format.space_before = Pt(24)
    style_h1.paragraph_format.space_after = Pt(6)
    rpr_h1 = style_h1.element.rPr
    rpr_h1.rFonts.set(qn('w:eastAsia'), font_name)

    # 标题2 (H2)
    style_h2 = document.styles.add_style('CustomHeading2', 1)
    style_h2.base_style = document.styles['Heading 2']
    font_h2 = style_h2.font
    font_h2.name = font_name
    font_h2.size = Pt(14)
    font_h2.bold = True
    font_h2.color.rgb = primary_color
    style_h2.paragraph_format.space_before = Pt(18)
    style_h2.paragraph_format.space_after = Pt(6)
    rpr_h2 = style_h2.element.rPr
    rpr_h2.rFonts.set(qn('w:eastAsia'), font_name)

    # 其他级别的标题可以类似添加...


def add_cover_page(document, report_title, subtitle, author, logo_path=None):
    """
    为文档添加一个专业的封面页。
    """
    # 添加Logo（如果提供）
    if logo_path and os.path.exists(logo_path):
        document.add_picture(logo_path, width=Inches(2.0))
        last_paragraph = document.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 主标题
    title_p = document.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(report_title)
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title_run.font.name = '微软雅黑'
    title_p.paragraph_format.space_before = Pt(72)
    title_p.paragraph_format.space_after = Pt(12)
    # 副标题
    subtitle_p = document.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_p.add_run(subtitle)
    subtitle_run.font.size = Pt(16)
    subtitle_run.font.italic = True
    subtitle_run.font.name = '微软雅黑'
    subtitle_p.paragraph_format.space_after = Pt(150)

    # 作者/公司信息
    author_p = document.add_paragraph()
    author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_run = author_p.add_run(f"评估方：{author}")
    author_run.font.size = Pt(12)
    author_run.font.name = '微软雅黑'

    # 日期
    from datetime import datetime
    date_p = document.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_p.add_run(f"生成日期：{datetime.now().strftime('%Y-%m-%d')}")
    date_run.font.size = Pt(12)
    date_run.font.name = '微软雅黑'
    document.add_page_break()


def add_header_footer(document,header_text,footer_text):
    """
    为文档的每个节添加页眉和页脚。
    """
    section = document.sections[0]
    # --- 页眉 ---
    header = section.header
    header_p = header.paragraphs[0]
    header_p.text = header_text
    # header_p.text = f"{company_name} - 简历评估报告"
    header_p.style = document.styles['Normal']
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # --- 页脚 ---
    footer = section.footer
    footer_p = footer.paragraphs[0]
    # 使用制表符将页码推到右侧
    footer_p.text = footer_text
    footer_p.style = document.styles['Normal']

    # 添加页码的复杂部分
    # 这是一个标准的添加页码字段的方法
    run = footer_p.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def style_table(table):
    """
    美化表格样式，包括表头背景色和字体。
    """
    table.style = 'Light Shading Accent 1'
    # 进一步自定义表头
    header_cells = table.rows[0].cells
    for cell in header_cells:
        # 设置背景色
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '4F81BD')  # 与标题颜色匹配的蓝色
        tc_pr.append(shd)

        # 设置表头字体为白色、加粗
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)

def set_custom_col_widths(document, table, wide_col_factor=4):
    """
    设置表格列宽，使最后两列比其他列宽。
    """
    num_cols = len(table.columns)
    num_wide_col = 3  # 设置有几个款列
    page_width = document.sections[0].page_width
    margins = document.sections[0].left_margin + document.sections[0].right_margin
    available_width = page_width - margins

    if num_cols < 3:
        if num_cols > 0:
            col_width = available_width / num_cols
            for col in table.columns:
                col.width = int(col_width)
        return

    total_weight = (num_cols - 2) * 1 + num_wide_col * wide_col_factor
    standard_col_width = available_width / total_weight
    wide_col_width = standard_col_width * wide_col_factor

    for i, column in enumerate(table.columns):
        if i < num_cols - 2:
            column.width = int(standard_col_width)
        else:
            column.width = int(wide_col_width)

def parse_inline_formatting(p, text: str, size=None):
    """
    解析内联格式（粗体、斜体、<br>换行）并将其添加到段落中。
    """
    p.clear()
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|<br\s*/?>)', text, flags=re.IGNORECASE)
    for part in filter(None, parts):
        run = None
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = p.add_run(part[1:-1])
            run.italic = True
        elif re.match(r'<br\s*/?>', part, re.IGNORECASE):
            p.add_run().add_break()
            continue
        else:
            run = p.add_run(part)

        if run and size:
            run.font.size = size


def is_table_row(line: str) -> bool:
    return line.strip().startswith('|') and line.strip().endswith('|')

def is_table_separator(line: str) -> bool:
    return bool(re.match(r'^\s*\|?(:?-+:?\|)+(:?-+:?)?\s*$', line.strip()))


def parse_markdown_to_docx(document, content: str):
    """
    解析包括表格在内的 Markdown 内容，并将其添加到文档中。
    表格的最后两列会比其他列更宽。
    """
    lines = content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]

        if not line.strip():
            i += 1
            continue

        # --- 标题解析逻辑 ---
        if line.startswith('#'):
            level = line.count('#')
            text = line.lstrip('# ').strip()
            heading_p = document.add_heading(level=min(level, 9))
            heading_p.clear()
            run = heading_p.add_run(text)
            font = run.font
            font.bold = True
            font.color.rgb = RGBColor(0, 0, 0)
            if level == 1:
                font.size = Pt(16)
            elif level == 2:
                font.size = Pt(14)
            else:
                font.size = Pt(13)
            i += 1
            continue

        # --- 表格解析逻辑 ---
        if (i + 1 < len(lines) and
                is_table_row(line) and
                is_table_separator(lines[i + 1])):

            header_line = line
            headers = [h.strip() for h in header_line.strip().strip('|').split('|')]
            num_cols = len(headers)

            table = document.add_table(rows=1, cols=num_cols)
            table.style = 'Table Grid'
            set_custom_col_widths(document, table)
            hdr_cells = table.rows[0].cells
            for j, header_text in enumerate(headers):
                p = hdr_cells[j].paragraphs[0]
                parse_inline_formatting(p, header_text, size=Pt(13))
                for run in p.runs: run.font.bold = True
                hdr_cells[j].vertical_alignment = 1

            i += 2
            while i < len(lines) and is_table_row(lines[i]):
                row_line = lines[i]
                row_data = [cell.strip() for cell in row_line.strip().strip('|').split('|')]
                row_cells = table.add_row().cells
                for j, cell_text in enumerate(row_data):
                    if j < num_cols:
                        p = row_cells[j].paragraphs[0]
                        parse_inline_formatting(p, cell_text, size=Pt(13))
                        row_cells[j].vertical_alignment = 1
                i += 1

            document.add_paragraph()
            continue

        # --- 其他元素解析逻辑 ---
        stripped_line = line.strip()
        if re.match(r'^[-\*]\s+', stripped_line):
            text = re.sub(r'^[-\*]\s+', '', stripped_line)
            p = document.add_paragraph(style='List Bullet')
            parse_inline_formatting(p, text, size=Pt(13))
        elif re.match(r'^\d+\.\s+', stripped_line):
            text = re.sub(r'^\d+\.\s+', '', stripped_line)
            p = document.add_paragraph(style='List Number')
            parse_inline_formatting(p, text, size=Pt(13))
        else:
            p = document.add_paragraph()
            parse_inline_formatting(p, line, size=Pt(13))

        i += 1

def create_resume_assessment_report(resume: Resume) -> dict:
    """
    根据简历内容生成一个包含Markdown格式文本和图片的、经过美化的Word评估报告。

    :param resume: 包含评估内容的字典。
    :return: 包含报告路径的字典。
    """
    # --- 1. 准备内容和路径 ---
    content_parts = [
        '# 1. 技术栈与岗位匹配度评价',
        resume.get('technology_stack_evaluate', '暂无技术栈评价。'),
        '# 2. 工作经历与岗位匹配度评价',
        resume.get('experience_evaluate', '暂无工作经历评价。'),
        '# 3. 简历书写结构化评价',
        resume.get('resume_struct_evaluate', '暂无简历结构评价。'),
    ]
    content = '\n\n'.join(content_parts)
    image_path = resume.get('resume_radar_path', '雷达图/five_dimension_radar.png')
    logo_path = 'logo.png' # <--- !!! 请替换为您的公司Logo路径
    company_name = "面面俱到"  # <--- !!! 请替换为您的公司名称
    output_dir = "简历评估"
    output_filename = os.path.join(output_dir, f"简历评估报告.docx")
    try:
        os.makedirs(output_dir, exist_ok=True)
    except OSError as e:
        print(f"创建目录 '{output_dir}' 时发生错误: {e}")
        return {'report_path': ''}

    # --- 2. 创建并美化文档 ---
    document = DocxDocument()
    # 步骤A: 设置全局样式
    setup_document_styles(document)
    # 步骤B: 添加封面
    add_cover_page(
        document,
        report_title="面试者综合能力评估报告",
        subtitle=f"针对面试者简历的深度分析",
        author=company_name,
        logo_path=logo_path
    )
    head_text= f"{company_name} - 简历评估报告"
    footer_text='一切输出仅供参考，请用户仔细甄别'
    # 步骤C: 添加页眉页脚
    add_header_footer(document, header_text=head_text, footer_text=footer_text)

    # --- 3. 解析并填充主要内容 ---
    parse_markdown_to_docx(document, content)

    # --- 4. 添加图片 ---
    document.add_heading('综合能力雷达图', level=2).style = 'CustomHeading2'
    if os.path.exists(image_path):
        try:
            # 添加图片并居中
            pic = document.add_picture(image_path, width=Inches(5.0))
            last_paragraph = document.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 为图片添加标题
            caption = document.add_paragraph(f'图1：综合能力雷达图')
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption.style = document.styles['Caption']

        except Exception as e:
            print(f"添加图片 '{image_path}' 时出错: {e}")
            document.add_paragraph(f"（无法加载图片：{image_path}）")
    else:
        print(f"警告：图片文件未找到：{image_path}。")
        document.add_paragraph(f"（图片文件丢失：{image_path}）")

    # --- 5. 保存文档 ---
    try:
        document.save(output_filename)
        print(f"报告已成功生成：{output_filename}")
        return {'report_path': output_filename}
    except Exception as e:
        print(f"保存Word文档时发生错误: {e}")
        return {'report_path': ''}


def generate_resume_pdf(content: str, image_path: str = "简历照片/微信图片_20250616170405.jpg"):
    """
    生成一个PDF格式的简历，包含Markdown格式的文本和一张图片。
    Args:
        content (str): Markdown格式的简历文本内容。
        image_path (str): 单个图片的完整路径（例如：'./profile.jpg'）。
    """
    output_pdf_path: str = "优化简历/resume.pdf"
    output_dir = os.path.dirname(output_pdf_path)
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir)

    if not os.path.exists(image_path):
        print(f"错误: 图片文件不存在于 '{image_path}'。请检查路径。")
        return
    image_uri = pathlib.Path(image_path).resolve().as_uri()
    html_content = markdown.markdown(content, extensions=['extra', 'nl2br'])
    # 定义CSS样式
    css_style = f"""
    @page {{
        size: A4;
        margin: 1in; /* A4页面的标准页边距 */
    }}
    body {{
        font-family: 'Segoe UI', 'Microsoft YaHei', sans-serif;
        line-height: 1.6;
        color: #333;
        font-size: 11pt;
    }}
    h1, h2, h3, h4, h5, h6 {{
        color: #2c3e50;
        margin-top: 1.5em;
        margin-bottom: 0.5em;
    }}
    h1 {{ font-size: 24pt; border-bottom: 2px solid #ddd; padding-bottom: 5px; }}
    h2 {{ font-size: 18pt; color: #34495e; }}
    h3 {{ font-size: 14pt; color: #555; }}

    /* ==================== 修改部分在这里 ==================== */
    .image-container {{
        position: absolute;
        /* 调整位置：减小top和right的值，让图片更靠近右上角 */
        top: 0.0in;
        right: 0.0in;
        /* 调整大小：将宽高从120px减小到100px */
        width: 90px;
        height: 90px;
        overflow: hidden;
        border-radius: 50%;
        border: 2px solid #eee;
        box-shadow: 0 4px 8px rgba(0,0,0,0.1);
        z-index: 1000;
    }}
    /* ======================================================== */

    .image-container img {{
        width: 100%;
        height: 100%;
        object-fit: cover;
        display: block;
    }}
    """
    full_html = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="utf-8">
        <title>个人简历</title>
        <style>{css_style}</style>
    </head>
    <body>
        <div class="image-container">
            <img src="{image_uri}" alt="个人头像">
        </div>
        {html_content}
    </body>
    </html>
    """
    try:
        HTML(string=full_html).write_pdf(output_pdf_path)
        print(f"简历PDF已成功生成至: {output_pdf_path}")
        return output_pdf_path
    except Exception as e:
        print(f"生成PDF时发生错误: {e}")
        print("请检查文件是否被其他程序占用，或检查文件夹权限。")

def create_interview_question_analyse_report(resume: Resume) -> dict:
    """
    根据简历内容生成一个包含Markdown格式文本和图片的、经过美化的Word评估报告。

    :param resume: 包含评估内容的字典。
    :return: 包含报告路径的字典。
    """
    # --- 1. 准备内容和路径 ---
    content_parts=['# 面试问题解析']
    i=1
    for question,standard_answer,answer,question_analyse,answer_analyse in zip(resume['question'],resume['standard_answer'],resume['answer'],resume['analyse'],resume['eval']):
        content_parts.append(
            f'##{i}. {question}\n'
            f'### 问题分析\n'
            f'{question_analyse}\n'
            f'### 范例答案\n'
            f'{standard_answer}\n'
            f'### 用户答案\n'
            f'{answer}\n'
            f'### 答案评估和改进建议\n'
            f'{answer_analyse}\n'
        )
        i+=1
    content = '\n\n'.join(content_parts)
    logo_path = 'logo.png' # <--- !!! 请替换为您的公司Logo路径
    company_name = "面面俱到"  # <--- !!! 请替换为您的公司名称
    output_dir = "问题解析"
    output_filename = os.path.join(output_dir, f"问题详解报告.docx")
    try:
        os.makedirs(output_dir, exist_ok=True)
    except OSError as e:
        print(f"创建目录 '{output_dir}' 时发生错误: {e}")
        return {'report_path': ''}

    # --- 2. 创建并美化文档 ---
    document = DocxDocument()
    # 步骤A: 设置全局样式
    setup_document_styles(document)
    # 步骤B: 添加封面
    add_cover_page(
        document,
        report_title="面试问题分析解答",
        subtitle=f"对面试者在面试中的回答和题目进行深度分析",
        author=company_name,
        logo_path=logo_path
    )
    head_text = f"{company_name} - 问题解析报告"
    footer_text = '一切输出仅供参考，请用户仔细甄别'
    # 步骤C: 添加页眉页脚
    add_header_footer(document, header_text=head_text, footer_text=footer_text)

    # --- 3. 解析并填充主要内容 ---
    parse_markdown_to_docx(document, content)
    # --- 5. 保存文档 ---
    try:
        document.save(output_filename)
        print(f"报告已成功生成：{output_filename}")
        return {'report_path': output_filename}
    except Exception as e:
        print(f"保存Word文档时发生错误: {e}")
        return {'report_path': ''}


